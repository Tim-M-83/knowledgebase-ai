import csv
import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.models.summarizer import SummarizerChunk, SummarizerDocument, SummarizerDocumentStatus
from app.services.chunking import RawSegment, chunk_segments
from app.services.embeddings import get_embeddings_provider
from app.services.summarizer_language import detect_language_code_from_texts
from app.utils.file_storage import load_file_path
from app.utils.text_cleaning import normalize_whitespace


logger = logging.getLogger(__name__)


def _extract_pdf(path: Path) -> list[RawSegment]:
    reader = PdfReader(str(path))
    segments: list[RawSegment] = []
    for idx, page in enumerate(reader.pages):
        text = normalize_whitespace(page.extract_text() or '')
        if text:
            segments.append(
                RawSegment(text=text, metadata={'source_type': 'pdf', 'page_number': idx + 1})
            )
    return segments


def _extract_txt(path: Path) -> list[RawSegment]:
    text = normalize_whitespace(path.read_text(encoding='utf-8', errors='ignore'))
    if not text:
        return []
    return [RawSegment(text=text, metadata={'source_type': 'txt'})]


def _extract_csv(path: Path) -> list[RawSegment]:
    segments: list[RawSegment] = []
    with path.open('r', encoding='utf-8', errors='ignore', newline='') as csvfile:
        reader = csv.reader(csvfile)
        rows = list(reader)

    for row_idx, row in enumerate(rows, start=1):
        row_text = normalize_whitespace(' | '.join(row))
        if row_text:
            segments.append(
                RawSegment(
                    text=row_text,
                    metadata={
                        'source_type': 'csv',
                        'csv_row_start': row_idx,
                        'csv_row_end': row_idx,
                    },
                )
            )
    return segments


def _extract_docx(path: Path) -> list[RawSegment]:
    doc = DocxDocument(str(path))
    segments: list[RawSegment] = []

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = normalize_whitespace(paragraph.text or '')
        if text:
            segments.append(
                RawSegment(
                    text=text,
                    metadata={'source_type': 'docx', 'paragraph_number': idx},
                )
            )

    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            row_values = [normalize_whitespace(cell.text or '') for cell in row.cells]
            row_text = normalize_whitespace(' | '.join([value for value in row_values if value]))
            if row_text:
                segments.append(
                    RawSegment(
                        text=row_text,
                        metadata={'source_type': 'docx', 'table_number': table_idx, 'row_number': row_idx},
                    )
                )

    return segments


def extract_summarizer_segments(document: SummarizerDocument) -> list[RawSegment]:
    path = load_file_path(document.filename)
    ext = document.original_name.split('.')[-1].lower()
    if ext == 'pdf':
        return _extract_pdf(path)
    if ext == 'txt':
        return _extract_txt(path)
    if ext == 'csv':
        return _extract_csv(path)
    if ext == 'docx':
        return _extract_docx(path)
    if ext == 'doc':
        raise ValueError('Legacy .doc files are not supported. Please upload DOCX.')
    raise ValueError(f'Unsupported file extension: {ext}')


def ingest_summarizer_document(db: Session, document_id: int) -> dict[str, Any]:
    document = db.query(SummarizerDocument).filter(SummarizerDocument.id == document_id).first()
    if not document:
        raise ValueError('Summarizer document not found')

    provider = get_embeddings_provider(db)

    try:
        logger.info('Summarizer ingestion started document_id=%s', document_id)
        document.status = SummarizerDocumentStatus.processing
        document.error_text = None
        db.commit()

        db.query(SummarizerChunk).filter(SummarizerChunk.document_id == document.id).delete()
        db.commit()

        segments = extract_summarizer_segments(document)
        if not segments:
            raise ValueError('No text could be extracted from this document')
        document.detected_language_code = detect_language_code_from_texts([segment.text for segment in segments])

        chunks = chunk_segments(segments)
        embeddings = provider.embed_batch([chunk.content for chunk in chunks])

        for chunk, embedding in zip(chunks, embeddings, strict=True):
            db.add(
                SummarizerChunk(
                    document_id=document.id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    embedding=embedding,
                    meta=chunk.metadata,
                )
            )

        document.status = SummarizerDocumentStatus.ready
        db.commit()
        logger.info('Summarizer ingestion succeeded document_id=%s chunks=%s', document.id, len(chunks))
        return {'chunks': len(chunks), 'document_id': document.id}
    except Exception as exc:
        db.rollback()
        document = db.query(SummarizerDocument).filter(SummarizerDocument.id == document_id).first()
        if document:
            document.status = SummarizerDocumentStatus.failed
            document.error_text = str(exc)
            db.commit()
        logger.exception('Summarizer ingestion failed document_id=%s: %s', document_id, exc)
        raise
